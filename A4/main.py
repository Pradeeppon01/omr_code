from docx import Document
from docx.shared import Inches,Pt
from barcode_gen import *
from PIL import Image
from csvReader import *
from constants import *
from doc_to_pdf import *
import time
import datetime
import shutil
import os


constant_qrno = 9826531965
# Define placeholders in the DOCX file and corresponding keys in the JSON data
placeholders = {
    "slno_placeholder": "SL_NO",
    "sno_placeholder":"SNO",
    "name_placeholder": "CANDIDATE_NAME",
    "reg_no_placeholder": "REGISTER_NUMBER",
    "department_placeholder": "department",
    "degree_placeholder": "degree",
    "degree_with_branch_placeholder": "DEGREE_WITH_BRANCH",
    "barcodeimg1": "barcode.png",
    "barcodeimg2": "barcodes/barcode2.gif",
    "qcode_placeholder":"QUESTION_CODE",
    "tabel_cell_1":"date",
    "exam_date_&_session_placeholder": "EXAM_DATE_&_SESSION",
    "exam_center_code": "EXAM_CENTER_CODE",
    "subject_code_placeholder":"SUBJECT_CODE",
    "subject_title_placeholder":"SUBJECT_TITLE",
    # "barcode1": "images/barcode1.gif",
    # "barcode2": "images/barcode2.gif",
    # "barcode_rot": "images/barcode3.gif",
    # "qrcode1": "images/qrcode.gif"
    "barcode1": "BARCODE1",
    "barcode2": "BARCODE2",
    "barcode3": "BARCODE3",
    "barcode4": "BARCODE4",
    "barcode5": "BARCODE5",
    "qrcode1": "QRCODE1"

    
}


def cleanupGeneratedBarcodes():
    dir1 = f"images/barcodes/"
    dir2 = f"images/rotated_imgs/"
    for filename in os.listdir(dir1):
        file_path = os.path.join(dir1, filename)
        os.remove(file_path)
    
    for filename in os.listdir(dir2):
        file_path = os.path.join(dir2, filename)
        os.remove(file_path)

def rotate_image(image_path, degrees):
    # Open the image
    image_name = os.path.basename(image_path)
    os.makedirs(img_output_dir, exist_ok=True)
    
    copied_output_path = os.path.join(img_output_dir, image_name)
    newimg_path = shutil.copy2(image_path, copied_output_path)
    print("newimg_path=",newimg_path)
    img = Image.open(newimg_path)
    # Rotate the image by the specified degrees
    rotated_img = img.rotate(degrees, expand=True)
    return rotated_img, newimg_path

def resizeImage(key):
    img = Image.open(key)
    # Resize the image to the desired dimensions
    desired_width_px = 300  # adjust as needed
    desired_height_px = 50  # adjust as needed
    img = img.resize((desired_width_px, desired_height_px), Image.ANTIALIAS)

    # Save the resized image with a temporary filename
    temp_img_path = "images/temp_resized_image.png"
    img.save(temp_img_path)
    return temp_img_path



def replace_placeholders_in_tables(tables, placeholders, json_data):
    # Iterate over tables
    for table in tables:
        # Iterate over rows
        for row in table.rows:
            # Iterate over cells in each row
            for cell in row.cells:
                # print("cell == ", cell.text)
                # Iterate over paragraphs in the cell
                for paragraph in cell.paragraphs:
                    # Iterate over placeholders
                    for placeholder, key in placeholders.items():
                        # Check if the placeholder is present in the paragraph text
                        if placeholder in paragraph.text:
                            # Replace the placeholder with corresponding JSON data
                            paragraph.text = paragraph.text.replace(placeholder, json_data.get(key, ""))
                            paragraph.style.font.size = Pt(8)
                            
                            # If the placeholder is for a barcode image, insert the image
                            if placeholder.startswith("barcode"):
                                barcode_val = json_data.get(key)
                                if placeholder == "barcode3" or placeholder == "barcode5" or placeholder == "barcode4":
                                    key = generateBarcode(barcode_val)
                                    rotated_image,newimg_path = rotate_image(key, -90)  # Rotate 90 degrees left
                                    rotated_image.save(newimg_path)  # Overwrite the original image with the rotated one
                                    paragraph.clear()
                                    # Add a run to the paragraph and insert the image
                                    run = paragraph.add_run()
                                    run.add_picture(newimg_path, width=Inches(0.55), height=Inches(1.44))  # Insert image
                                
                                else:
                                    key = generateBarcode(barcode_val)
                                    # img_path = resizeImage(key)
                                    # print("img_path ================================ ",img_path)
                                    # time.sleep(5)
                                    # Clear the paragraph to remove existing text
                                    paragraph.clear()
                                    # Add a run to the paragraph and insert the image
                                    run = paragraph.add_run()
                                    run.add_picture(key, width=Inches(1.49), height=Inches(0.35))  # Insert image
                                    
                            elif placeholder.startswith("qrcode"):
                                qrval = json_data.get(key)
                                print("barcode_val = ", qrval)
                                key = generateQRCode(qrval)
                                paragraph.clear()
                                # Add a run to the paragraph and insert the image
                                run = paragraph.add_run()
                                run.add_picture(key, width=Inches(0.35), height=Inches(0.35))  # Insert image
                # Check for nested tables
                nested_tables = cell.tables
                # print("nested_tables ==",nested_tables)
                if nested_tables:
                    # Recursively replace placeholders in nested tables
                    replace_placeholders_in_tables(nested_tables, placeholders, json_data)



# Save the modified DOCX file (overwrite the original file)
def entry_point(doc, json_data):
    print("Execution started....")
    # Replace placeholders in main document content
    for paragraph in doc.paragraphs:
        print("paragraph === ", paragraph.text)
        # Iterate over placeholders
        for placeholder, key in placeholders.items():
            print("................",placeholder,key)
            # Check if the placeholder is present in the paragraph text
            if placeholder in paragraph.text:
                # Replace the placeholder with corresponding JSON data
                paragraph.text = paragraph.text.replace(placeholder, json_data.get(key, ""))
                paragraph.style.font.size = Pt(8)
                # If the placeholder is for a barcode image, insert the image
                                            # If the placeholder is for a barcode image, insert the image
                if placeholder.startswith("barcode"):
                    barcode_val = json_data.get(key)
                    print("barcode_val = ", barcode_val)
                    if placeholder == "barcode3" or placeholder == "barcode5" or placeholder == "barcode4":
                        #generate barcode here and store it in images/barcodes/
                        key = generateBarcode(barcode_val)
                        rotated_image,newimg_path = rotate_image(key, -90)  # Rotate 90 degrees left
                        rotated_image.save(newimg_path)  # Overwrite the original image with the rotated one
                        paragraph.clear()
                        # Add a run to the paragraph and insert the image
                        run = paragraph.add_run()
                        run.add_picture(newimg_path, width=Inches(0.47), height=Inches(1))  # Insert image
                    
                    else:
                        #generate barcode here and store it in images/barcodes/
                        key = generateBarcode(barcode_val)
                        # img_path = resizeImage(key)
                        # print("img_path ================================ ",img_path)
                        # time.sleep(5)
                        # Clear the paragraph to remove existing text
                        paragraph.clear()
                        # Add a run to the paragraph and insert the image
                        run = paragraph.add_run()
                        run.add_picture(key, width=Inches(1.49), height=Inches(0.35))  # Insert image
                
                elif placeholder.startswith("qrcode"):
                    qrval = json_data.get(key)
                    print("barcode_val = ", qrval)
                    key = generateQRCode(qrval)
                    paragraph.clear()
                    # Add a run to the paragraph and insert the image
                    run = paragraph.add_run()
                    run.add_picture(key, width=Inches(0.35), height=Inches(0.35))  # Insert image

    # Replace placeholders in tables
    replace_placeholders_in_tables(doc.tables, placeholders, json_data)
    


if __name__ == "__main__":
    csv_file = 'assets/studentData.csv'
    start_time = time.time()
    data_dict = read_csv_to_dict(csv_file)
    #print("data_dict = ", data_dict)
    for obj in data_dict:
        print("obj === ", data_dict[obj])
        json_data = data_dict[obj]
        print("json_data = ",json_data)

        img_output_dir = "images/rotated_imgs/"
        # Open the DOCX file
        # doc = Document("parta-a41.docx")
        doc = Document("full_a3sheet.docx")
        entry_point(doc, json_data)
        
        doc.save(f"output_doc/{obj}.docx")
        cleanupGeneratedBarcodes()
        
        #break
    pdfConversionMain()
    end_time = time.time()
    elapsed_time = end_time - start_time
    print("Execution time: {:.2f} seconds".format(elapsed_time))
        
